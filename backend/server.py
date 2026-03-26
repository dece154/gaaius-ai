from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form, Depends, Request
from fastapi.responses import StreamingResponse, FileResponse, HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import asyncio
import base64
import io
import hashlib
import jwt
import requests
import json

# Load environment variables
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# API Keys
GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
HF_TOKEN = os.environ.get('HF_TOKEN')
PAYPAL_CLIENT_ID = os.environ.get('PAYPAL_CLIENT_ID')
PAYPAL_SECRET = os.environ.get('PAYPAL_SECRET')
PAYFAST_MERCHANT_ID = os.environ.get('PAYFAST_MERCHANT_ID')
PAYFAST_MERCHANT_KEY = os.environ.get('PAYFAST_MERCHANT_KEY')
JWT_SECRET = os.environ.get('JWT_SECRET', 'default_secret')

# Initialize Groq client
from groq import Groq
groq_client = Groq(api_key=GROQ_API_KEY)

# Initialize Hugging Face client
from huggingface_hub import InferenceClient
hf_client = InferenceClient(api_key=HF_TOKEN)

# Create the main app
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer(auto_error=False)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ============== MODELS ==============

class UserRegister(BaseModel):
    email: str
    password: str
    name: str = ""

class UserLogin(BaseModel):
    email: str
    password: str

class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str = ""
    password_hash: str
    is_pro: bool = False
    pro_expires: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ChatMessage(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    role: str
    content: str
    model_used: Optional[str] = None
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatRequest(BaseModel):
    session_id: str
    message: str
    
class ChatResponse(BaseModel):
    id: str
    content: str
    model_used: str
    timestamp: str

class ImageGenerationRequest(BaseModel):
    prompt: str
    session_id: Optional[str] = None

class ImageGenerationResponse(BaseModel):
    id: str
    prompt: str
    image_url: str
    model_used: str
    timestamp: str

class VideoGenerationRequest(BaseModel):
    prompt: str
    duration: int = 5
    style: str = "cinematic"
    session_id: Optional[str] = None

class VideoGenerationResponse(BaseModel):
    id: str
    prompt: str
    video_url: str
    model_used: str
    timestamp: str

class TTSRequest(BaseModel):
    text: str
    voice: str = "en"

class AudioGenerationRequest(BaseModel):
    prompt: str
    duration: int = 10
    type: str = "music"  # music, sfx, ambient
    voice: str = "default"  # default, male, female
    language: str = ""  # empty = auto-detect

class FileGenerationRequest(BaseModel):
    prompt: str
    file_type: str  # code, document, data, config

class ProjectCreate(BaseModel):
    name: str
    description: str = ""
    type: str = "web"  # web, api, data

class Session(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str = "New Chat"
    user_id: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# ============== AUTH HELPERS ==============

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def create_token(user_id: str, email: str, is_pro: bool) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "is_pro": is_pro,
        "exp": datetime.now(timezone.utc) + timedelta(days=30)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if not credentials:
        return None
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=["HS256"])
        user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
        return user
    except:
        return None

# ============== AUTH ROUTES ==============

@api_router.post("/auth/register")
async def register(data: UserRegister):
    existing = await db.users.find_one({"email": data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user = User(
        email=data.email,
        name=data.name,
        password_hash=hash_password(data.password)
    )
    await db.users.insert_one(user.model_dump())
    token = create_token(user.id, user.email, user.is_pro)
    
    return {"token": token, "user": {"id": user.id, "email": user.email, "name": user.name, "is_pro": user.is_pro}}

@api_router.post("/auth/login")
async def login(data: UserLogin):
    user = await db.users.find_one({"email": data.email}, {"_id": 0})
    if not user or user["password_hash"] != hash_password(data.password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check pro status
    is_pro = user.get("is_pro", False)
    if is_pro and user.get("pro_expires"):
        if datetime.fromisoformat(user["pro_expires"]) < datetime.now(timezone.utc):
            is_pro = False
            await db.users.update_one({"id": user["id"]}, {"$set": {"is_pro": False}})
    
    token = create_token(user["id"], user["email"], is_pro)
    return {"token": token, "user": {"id": user["id"], "email": user["email"], "name": user.get("name", ""), "is_pro": is_pro}}

@api_router.get("/auth/me")
async def get_me(user = Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return {"id": user["id"], "email": user["email"], "name": user.get("name", ""), "is_pro": user.get("is_pro", False)}

# ============== PAYMENT ROUTES ==============

@api_router.post("/payment/paypal/create")
async def create_paypal_order(user = Depends(get_current_user)):
    """Create PayPal order for Pro subscription"""
    try:
        # Get PayPal access token
        auth = base64.b64encode(f"{PAYPAL_CLIENT_ID}:{PAYPAL_SECRET}".encode()).decode()
        token_response = requests.post(
            "https://api-m.paypal.com/v1/oauth2/token",
            headers={"Authorization": f"Basic {auth}", "Content-Type": "application/x-www-form-urlencoded"},
            data="grant_type=client_credentials"
        )
        access_token = token_response.json().get("access_token")
        
        # Create order
        order_response = requests.post(
            "https://api-m.paypal.com/v2/checkout/orders",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json"
            },
            json={
                "intent": "CAPTURE",
                "purchase_units": [{
                    "amount": {"currency_code": "USD", "value": "1.00"},
                    "description": "GAAIUS AI Pro - 1 Month"
                }]
            }
        )
        return order_response.json()
    except Exception as e:
        logger.error(f"PayPal create error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/payment/paypal/capture/{order_id}")
async def capture_paypal_order(order_id: str, user = Depends(get_current_user)):
    """Capture PayPal payment and activate Pro"""
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        auth = base64.b64encode(f"{PAYPAL_CLIENT_ID}:{PAYPAL_SECRET}".encode()).decode()
        token_response = requests.post(
            "https://api-m.paypal.com/v1/oauth2/token",
            headers={"Authorization": f"Basic {auth}", "Content-Type": "application/x-www-form-urlencoded"},
            data="grant_type=client_credentials"
        )
        access_token = token_response.json().get("access_token")
        
        capture_response = requests.post(
            f"https://api-m.paypal.com/v2/checkout/orders/{order_id}/capture",
            headers={"Authorization": f"Bearer {access_token}", "Content-Type": "application/json"}
        )
        result = capture_response.json()
        
        if result.get("status") == "COMPLETED":
            # Activate Pro for 30 days
            expires = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
            await db.users.update_one(
                {"id": user["id"]},
                {"$set": {"is_pro": True, "pro_expires": expires}}
            )
            await db.payments.insert_one({
                "id": str(uuid.uuid4()),
                "user_id": user["id"],
                "provider": "paypal",
                "order_id": order_id,
                "amount": 1.00,
                "currency": "USD",
                "status": "completed",
                "timestamp": datetime.now(timezone.utc).isoformat()
            })
            return {"success": True, "message": "Pro activated!", "expires": expires}
        
        raise HTTPException(status_code=400, detail="Payment not completed")
    except Exception as e:
        logger.error(f"PayPal capture error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/payment/payfast/create")
async def create_payfast_payment(user = Depends(get_current_user)):
    """Generate PayFast payment URL"""
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    payment_id = str(uuid.uuid4())
    
    # PayFast payment data
    data = {
        "merchant_id": PAYFAST_MERCHANT_ID,
        "merchant_key": PAYFAST_MERCHANT_KEY,
        "return_url": f"https://groq-chat-preview.preview.emergentagent.com/?payment=success&id={payment_id}",
        "cancel_url": "https://groq-chat-preview.preview.emergentagent.com/?payment=cancelled",
        "notify_url": f"https://groq-chat-preview.preview.emergentagent.com/api/payment/payfast/notify",
        "amount": "18.00",  # ~$1 in ZAR
        "item_name": "GAAIUS AI Pro - 1 Month",
        "custom_str1": user["id"],
        "custom_str2": payment_id
    }
    
    # Generate signature
    param_string = "&".join([f"{k}={v}" for k, v in sorted(data.items()) if k != "signature"])
    signature = hashlib.md5(param_string.encode()).hexdigest()
    data["signature"] = signature
    
    # Store pending payment
    await db.payments.insert_one({
        "id": payment_id,
        "user_id": user["id"],
        "provider": "payfast",
        "amount": 18.00,
        "currency": "ZAR",
        "status": "pending",
        "timestamp": datetime.now(timezone.utc).isoformat()
    })
    
    return {"payment_url": "https://www.payfast.co.za/eng/process", "data": data}

@api_router.post("/payment/payfast/notify")
async def payfast_notify(request: Request):
    """PayFast ITN callback"""
    try:
        form_data = await request.form()
        data = dict(form_data)
        
        if data.get("payment_status") == "COMPLETE":
            user_id = data.get("custom_str1")
            payment_id = data.get("custom_str2")
            
            expires = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
            await db.users.update_one({"id": user_id}, {"$set": {"is_pro": True, "pro_expires": expires}})
            await db.payments.update_one({"id": payment_id}, {"$set": {"status": "completed"}})
        
        return {"status": "ok"}
    except Exception as e:
        logger.error(f"PayFast notify error: {e}")
        return {"status": "error"}

@api_router.get("/payment/config")
async def get_payment_config():
    """Get payment configuration for frontend"""
    return {
        "paypal_client_id": PAYPAL_CLIENT_ID,
        "payfast_merchant_id": PAYFAST_MERCHANT_ID,
        "pro_price_usd": 1.00,
        "pro_price_zar": 18.00
    }

# ============== BASIC ROUTES ==============

@api_router.get("/")
async def root():
    return {"message": "GAAIUS AI Backend Running", "status": "operational"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "groq": bool(GROQ_API_KEY), "huggingface": bool(HF_TOKEN)}

# ============== SESSION ROUTES ==============

@api_router.post("/sessions", response_model=dict)
async def create_session(name: str = "New Chat", user = Depends(get_current_user)):
    session = Session(name=name, user_id=user["id"] if user else None)
    doc = session.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['updated_at'] = doc['updated_at'].isoformat()
    await db.sessions.insert_one(doc)
    return {"id": session.id, "name": session.name, "created_at": doc['created_at']}

@api_router.put("/sessions/{session_id}")
async def update_session(session_id: str, data: dict):
    """Update session name"""
    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if "name" in data:
        update_data["name"] = data["name"]
    await db.sessions.update_one({"id": session_id}, {"$set": update_data})
    return {"status": "updated"}

@api_router.get("/sessions")
async def get_sessions(user = Depends(get_current_user)):
    query = {"user_id": user["id"]} if user else {}
    sessions = await db.sessions.find(query, {"_id": 0}).sort("updated_at", -1).to_list(100)
    return sessions

@api_router.delete("/sessions/{session_id}")
async def delete_session(session_id: str):
    await db.sessions.delete_one({"id": session_id})
    await db.messages.delete_many({"session_id": session_id})
    return {"status": "deleted"}

# ============== CHAT ROUTES ==============

@api_router.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, user = Depends(get_current_user)):
    try:
        history = await db.messages.find({"session_id": request.session_id}, {"_id": 0}).sort("timestamp", 1).to_list(50)
        
        # Get current date and time for real-time awareness
        from datetime import datetime
        import pytz
        now = datetime.now(pytz.UTC)
        current_date = now.strftime("%A, %B %d, %Y")
        current_time = now.strftime("%I:%M %p UTC")
        
        system_prompt = f"""You are GAAIUS AI, a powerful, friendly, and intelligent AI assistant similar to ChatGPT. 

**Current Date & Time:** {current_date}, {current_time}

**Your Capabilities:**
- Engage in natural, helpful conversations
- Answer questions with accurate, well-structured responses
- Help with coding, writing, analysis, math, and creative tasks
- Provide thoughtful explanations with examples when helpful
- Remember context from the conversation

**Response Style:**
- Be conversational, warm, and helpful like ChatGPT
- Use clear formatting with paragraphs for readability
- Use bullet points or numbered lists when listing items
- Use code blocks with syntax highlighting for code
- Provide comprehensive but concise answers
- Ask clarifying questions when needed
- Be honest when you don't know something

**Important:** You have access to the current date and time above. Use it when users ask about today's date, current time, or time-sensitive questions."""
        
        messages = [{"role": "system", "content": system_prompt}]
        
        for msg in history:
            messages.append({"role": msg['role'], "content": msg['content']})
        
        messages.append({"role": "user", "content": request.message})
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.7,
            max_tokens=2048
        )
        
        response_content = completion.choices[0].message.content
        model_used = "Groq Llama 3.3 70B"
        
        # Save messages
        user_msg = ChatMessage(session_id=request.session_id, role="user", content=request.message)
        user_doc = user_msg.model_dump()
        user_doc['timestamp'] = user_doc['timestamp'].isoformat()
        await db.messages.insert_one(user_doc)
        
        assistant_msg = ChatMessage(session_id=request.session_id, role="assistant", content=response_content, model_used=model_used)
        assistant_doc = assistant_msg.model_dump()
        assistant_doc['timestamp'] = assistant_doc['timestamp'].isoformat()
        await db.messages.insert_one(assistant_doc)
        
        await db.sessions.update_one({"id": request.session_id}, {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}})
        
        return ChatResponse(id=assistant_msg.id, content=response_content, model_used=model_used, timestamp=assistant_doc['timestamp'])
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/tts/speak")
async def text_to_speech(data: dict):
    """Convert text to speech in any language"""
    try:
        from gtts import gTTS
        
        text = data.get("text", "")
        lang = data.get("lang", "en")  # Default English
        
        if not text:
            raise HTTPException(status_code=400, detail="Text is required")
        
        # Supported languages
        lang_map = {
            "en": "en", "english": "en",
            "es": "es", "spanish": "es",
            "fr": "fr", "french": "fr",
            "de": "de", "german": "de",
            "it": "it", "italian": "it",
            "pt": "pt", "portuguese": "pt",
            "zh": "zh-CN", "chinese": "zh-CN",
            "ja": "ja", "japanese": "ja",
            "ko": "ko", "korean": "ko",
            "ru": "ru", "russian": "ru",
            "ar": "ar", "arabic": "ar",
            "hi": "hi", "hindi": "hi",
            "nl": "nl", "dutch": "nl",
            "pl": "pl", "polish": "pl",
            "tr": "tr", "turkish": "tr",
            "vi": "vi", "vietnamese": "vi",
            "th": "th", "thai": "th",
            "id": "id", "indonesian": "id",
            "sv": "sv", "swedish": "sv",
            "da": "da", "danish": "da",
            "no": "no", "norwegian": "no",
            "fi": "fi", "finnish": "fi",
            "el": "el", "greek": "el",
            "he": "he", "hebrew": "he",
            "cs": "cs", "czech": "cs",
            "ro": "ro", "romanian": "ro",
            "hu": "hu", "hungarian": "hu",
            "uk": "uk", "ukrainian": "uk",
            "af": "af", "afrikaans": "af",
            "zu": "zu", "zulu": "zu",
            "sw": "sw", "swahili": "sw"
        }
        
        tts_lang = lang_map.get(lang.lower(), "en")
        
        gen_id = str(uuid.uuid4())
        audio_filename = f"tts_{gen_id}.mp3"
        audio_path = ROOT_DIR / "static" / "audio" / audio_filename
        (ROOT_DIR / "static" / "audio").mkdir(parents=True, exist_ok=True)
        
        tts = gTTS(text=text, lang=tts_lang, slow=False)
        tts.save(str(audio_path))
        
        audio_url = f"/api/static/audio/{audio_filename}"
        
        return {"audio_url": audio_url, "language": tts_lang}
        
    except Exception as e:
        logger.error(f"TTS error: {e}")
        raise HTTPException(status_code=500, detail=f"TTS failed: {str(e)}")

@api_router.get("/chat/{session_id}/history")
async def get_chat_history(session_id: str):
    messages = await db.messages.find({"session_id": session_id}, {"_id": 0}).sort("timestamp", 1).to_list(1000)
    return messages

# ============== IMAGE GENERATION ==============

@api_router.post("/image/generate", response_model=ImageGenerationResponse)
async def generate_image(request: ImageGenerationRequest, user = Depends(get_current_user)):
    try:
        import requests as req
        from PIL import Image as PILImage
        import urllib.parse
        
        gen_id = str(uuid.uuid4())
        img_filename = f"{gen_id}.jpg"
        img_path = ROOT_DIR / "static" / img_filename
        (ROOT_DIR / "static").mkdir(exist_ok=True)
        
        image_bytes = None
        model_used = "Unknown"
        
        # Pollinations.ai with retry (3 attempts, increasing timeout)
        encoded_prompt = urllib.parse.quote(request.prompt)
        pollinations_urls = [
            f"https://image.pollinations.ai/prompt/{encoded_prompt}?nologo=true&width=1024&height=1024",
            f"https://image.pollinations.ai/prompt/{encoded_prompt}?nologo=true",
            f"https://image.pollinations.ai/prompt/{encoded_prompt}",
        ]
        
        for attempt, url in enumerate(pollinations_urls):
            try:
                timeout = 30 + (attempt * 30)  # 30s, 60s, 90s
                logger.info(f"Pollinations attempt {attempt+1}: timeout={timeout}s")
                response = req.get(url, timeout=timeout)
                if response.status_code == 200 and len(response.content) > 5000:
                    image_bytes = response.content
                    model_used = "Pollinations AI"
                    logger.info(f"Pollinations succeeded on attempt {attempt+1}")
                    break
            except Exception as e:
                logger.warning(f"Pollinations attempt {attempt+1} failed: {e}")
                continue
        
        # Fallback: HuggingFace (only if token is valid)
        if not image_bytes and HF_TOKEN:
            for hf_model in ["stabilityai/stable-diffusion-xl-base-1.0", "black-forest-labs/FLUX.1-dev"]:
                try:
                    logger.info(f"Trying HuggingFace: {hf_model}")
                    hf_image = hf_client.text_to_image(request.prompt, model=hf_model)
                    if hf_image:
                        buf = io.BytesIO()
                        hf_image.save(buf, format='JPEG', quality=90)
                        image_bytes = buf.getvalue()
                        model_used = f"HuggingFace {hf_model.split('/')[-1]}"
                        break
                except Exception as e:
                    logger.warning(f"HuggingFace {hf_model} failed: {e}")
        
        if not image_bytes:
            raise Exception("Image generation is temporarily unavailable. Please try again in a moment.")
        
        image = PILImage.open(io.BytesIO(image_bytes))
        image = image.convert("RGB")
        image.save(img_path, format='JPEG', quality=90)
        
        image_url = f"/api/static/{img_filename}"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "image", "prompt": request.prompt, "url": image_url,
            "model_used": model_used, "session_id": request.session_id, "timestamp": timestamp
        })
        
        return ImageGenerationResponse(id=gen_id, prompt=request.prompt, image_url=image_url, model_used=model_used, timestamp=timestamp)
        
    except Exception as e:
        logger.error(f"Image generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Image generation failed: {str(e)}")

# ============== VIDEO GENERATION ==============

from video_engine import VideoEngine, StoryVideoEngine

video_engine = VideoEngine(hf_token=HF_TOKEN, groq_api_key=GROQ_API_KEY, output_dir=ROOT_DIR / "static" / "videos")
story_video_engine = StoryVideoEngine(hf_token=HF_TOKEN, groq_api_key=GROQ_API_KEY, output_dir=ROOT_DIR / "static" / "videos")

@api_router.post("/video/generate", response_model=VideoGenerationResponse)
async def generate_video(request: VideoGenerationRequest, user = Depends(get_current_user)):
    try:
        result = await video_engine.generate_video(
            prompt=request.prompt, duration=min(request.duration, 30), fps=8, style=request.style
        )
        
        video_filename = Path(result["video_path"]).name
        video_url = f"/api/static/videos/{video_filename}"
        model_used = f"GAAIUS Video Engine ({request.style})"
        gen_id = result["video_id"]
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "video", "prompt": request.prompt, "url": video_url,
            "model_used": model_used, "session_id": request.session_id, "timestamp": timestamp
        })
        
        return VideoGenerationResponse(id=gen_id, prompt=request.prompt, video_url=video_url, model_used=model_used, timestamp=timestamp)
        
    except Exception as e:
        logger.error(f"Video generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Video generation failed: {str(e)}")

@api_router.post("/video/generate-story")
async def generate_story_video(request: dict, user = Depends(get_current_user)):
    try:
        result = await story_video_engine.generate_story_video(
            story_prompt=request.get("prompt", ""),
            chapters=min(request.get("chapters", 3), 5),
            duration_per_chapter=min(request.get("duration_per_chapter", 8), 15),
            style=request.get("style", "cinematic")
        )
        
        video_filename = Path(result["video_path"]).name
        video_url = f"/api/static/videos/{video_filename}"
        
        return {"id": result["video_id"], "video_url": video_url, "chapters": result.get("chapters", []), "total_duration": result.get("total_duration", 0)}
    except Exception as e:
        logger.error(f"Story video error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ============== AUDIO GENERATION (HuggingFace TTS/STT) ==============

@api_router.post("/tts")
async def text_to_speech(request: TTSRequest, user = Depends(get_current_user)):
    """Text-to-Speech using HuggingFace"""
    try:
        # Try multiple TTS models as fallback
        models_to_try = [
            "espnet/kan-bayashi_ljspeech_vits",
            "facebook/mms-tts-eng",
            "microsoft/speecht5_tts"
        ]
        
        audio = None
        last_error = None
        
        for model in models_to_try:
            try:
                audio = hf_client.text_to_speech(request.text, model=model)
                if audio:
                    break
            except Exception as e:
                last_error = e
                continue
        
        if not audio:
            raise last_error or Exception("TTS failed with all models")
        
        return StreamingResponse(io.BytesIO(audio), media_type="audio/wav", headers={"Content-Disposition": "attachment; filename=speech.wav"})
        
    except Exception as e:
        logger.error(f"TTS error: {e}")
        raise HTTPException(status_code=500, detail=f"TTS failed: {str(e)}")

@api_router.post("/stt")
async def speech_to_text(audio: UploadFile = File(...), user = Depends(get_current_user)):
    """Speech-to-Text using HuggingFace Whisper"""
    try:
        audio_content = await audio.read()
        
        # Use Whisper via HuggingFace
        result = hf_client.automatic_speech_recognition(audio_content, model="openai/whisper-large-v3")
        
        text = result.get("text", "") if isinstance(result, dict) else str(result)
        return {"text": text, "model_used": "Whisper (HuggingFace)"}
        
    except Exception as e:
        logger.error(f"STT error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/audio/generate")
async def generate_audio(request: AudioGenerationRequest, user = Depends(get_current_user)):
    """Generate audio from any prompt - stories, narration, anything. Supports voice and language selection."""
    try:
        from gtts import gTTS
        import re
        
        gen_id = str(uuid.uuid4())
        audio_filename = f"{gen_id}.mp3"
        audio_path = ROOT_DIR / "static" / "audio" / audio_filename
        (ROOT_DIR / "static" / "audio").mkdir(parents=True, exist_ok=True)
        
        # Get voice/language from request
        voice_type = request.voice or 'default'
        lang = request.language if request.language else None
        
        if not lang:
            # Detect language from prompt
            prompt_lower = request.prompt.lower()
            lang = 'en'  # Default English
            
            lang_keywords = {
                'es': ['spanish', 'español', 'espanol'],
                'fr': ['french', 'français', 'francais'],
                'de': ['german', 'deutsch'],
                'it': ['italian', 'italiano'],
                'pt': ['portuguese', 'português'],
                'zh-CN': ['chinese', '中文', 'mandarin'],
                'ja': ['japanese', '日本語'],
                'ko': ['korean', '한국어'],
                'ru': ['russian', 'русский'],
                'ar': ['arabic', 'عربي'],
                'hi': ['hindi', 'हिंदी'],
                'af': ['afrikaans'],
                'zu': ['zulu'],
                'sw': ['swahili'],
                'nl': ['dutch'],
                'pl': ['polish'],
                'tr': ['turkish'],
                'vi': ['vietnamese'],
                'th': ['thai'],
                'id': ['indonesian']
            }
            
            for code, keywords in lang_keywords.items():
                if any(kw in prompt_lower for kw in keywords):
                    lang = code
                    break
        
        # Extract duration if specified
        prompt_lower = request.prompt.lower()
        duration_match = re.search(r'(\d+)\s*(minute|min|second|sec)', prompt_lower)
        target_words = 150  # Default ~1 minute
        if duration_match:
            num = int(duration_match.group(1))
            unit = duration_match.group(2)
            if 'min' in unit:
                target_words = num * 150
            else:
                target_words = max(30, num * 2)
        
        # Use AI to create content from the prompt
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": f"""You are a professional narrator and storyteller. Create exactly what the user asks for.

If they ask for a story, write a complete engaging story.
If they ask for narration, narrate it professionally.
If they just type something, read it naturally as narration.

Guidelines:
- Target approximately {target_words} words
- Write naturally for audio (no emojis, no formatting symbols)
- Make it engaging and interesting
- If a duration is specified, match it (150 words ≈ 1 minute)"""},
                {"role": "user", "content": request.prompt}
            ],
            temperature=0.8,
            max_tokens=min(4000, target_words * 2)
        )
        narration_text = completion.choices[0].message.content
        
        # Convert to speech
        tts = gTTS(text=narration_text, lang=lang, slow=False)
        tts.save(str(audio_path))
        
        audio_url = f"/api/static/audio/{audio_filename}"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        lang_names = {
            'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German', 
            'it': 'Italian', 'pt': 'Portuguese', 'zh-CN': 'Chinese', 'ja': 'Japanese',
            'ko': 'Korean', 'ru': 'Russian', 'ar': 'Arabic', 'hi': 'Hindi',
            'af': 'Afrikaans', 'zu': 'Zulu', 'sw': 'Swahili', 'nl': 'Dutch',
            'pl': 'Polish', 'tr': 'Turkish', 'vi': 'Vietnamese', 'th': 'Thai', 'id': 'Indonesian'
        }
        
        await db.generations.insert_one({
            "id": gen_id, "type": "audio", "prompt": request.prompt, "url": audio_url,
            "content": narration_text, "language": lang_names.get(lang, 'English'), 
            "voice": voice_type, "timestamp": timestamp
        })
        
        return {
            "id": gen_id, 
            "audio_url": audio_url, 
            "content": narration_text, 
            "language": lang_names.get(lang, 'English'),
            "voice": voice_type,
            "timestamp": timestamp
        }
        
    except Exception as e:
        logger.error(f"Audio generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Audio generation failed: {str(e)}")

# ============== FILE GENERATION ==============

@api_router.post("/file/generate")
async def generate_file(request: FileGenerationRequest, user = Depends(get_current_user)):
    """Generate code/documents including PDF, DOCX, XLSX"""
    try:
        system_prompts = {
            "code": "You are an expert programmer. Generate clean, well-documented code based on the user's request. Output only the code, no explanations.",
            "document": "You are a professional document writer. Generate well-structured content with clear sections and paragraphs. Use markdown formatting with # for headers.",
            "data": "You are a data expert. Generate sample data in the exact format requested - JSON, CSV, XML. Output only the data.",
            "config": "You are a DevOps expert. Generate configuration files. Output only the config, no explanations."
        }
        
        prompt_lower = request.prompt.lower()
        ext = "txt"
        is_binary = False
        
        # Detect file format
        if "pdf" in prompt_lower:
            ext = "pdf"
            is_binary = True
        elif "docx" in prompt_lower or "word" in prompt_lower:
            ext = "docx"
            is_binary = True
        elif "xlsx" in prompt_lower or "excel" in prompt_lower:
            ext = "xlsx"
            is_binary = True
        elif request.file_type == "code":
            if "python" in prompt_lower or ".py" in prompt_lower:
                ext = "py"
            elif "javascript" in prompt_lower or ".js" in prompt_lower:
                ext = "js"
            elif "typescript" in prompt_lower or ".ts" in prompt_lower:
                ext = "ts"
            elif "html" in prompt_lower:
                ext = "html"
            elif "css" in prompt_lower:
                ext = "css"
            else:
                ext = "py"
        elif request.file_type == "document":
            if "html" in prompt_lower:
                ext = "html"
            elif "txt" in prompt_lower:
                ext = "txt"
            else:
                ext = "md"
        elif request.file_type == "data":
            if "csv" in prompt_lower:
                ext = "csv"
            elif "xml" in prompt_lower:
                ext = "xml"
            else:
                ext = "json"
        elif request.file_type == "config":
            if "yaml" in prompt_lower or "yml" in prompt_lower:
                ext = "yaml"
            elif "toml" in prompt_lower:
                ext = "toml"
            else:
                ext = "json"
        
        # Generate content using Groq
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompts.get(request.file_type, system_prompts["document"])},
                {"role": "user", "content": request.prompt}
            ],
            temperature=0.3,
            max_tokens=4096
        )
        
        content = completion.choices[0].message.content
        
        # Clean up code blocks
        if "```" in content:
            import re
            code_match = re.search(r'```[\w]*\n?([\s\S]*?)```', content)
            if code_match:
                content = code_match.group(1).strip()
        
        gen_id = str(uuid.uuid4())
        file_filename = f"{gen_id}.{ext}"
        file_path = ROOT_DIR / "static" / "files" / file_filename
        (ROOT_DIR / "static" / "files").mkdir(parents=True, exist_ok=True)
        
        # Generate binary files (PDF, DOCX, XLSX)
        if ext == "pdf":
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            
            doc = SimpleDocTemplate(str(file_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            for line in content.split('\n'):
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
            
            doc.build(story)
            
        elif ext == "docx":
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
            doc.save(str(file_path))
            
        elif ext == "xlsx":
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Data"
            
            for i, line in enumerate(content.split('\n'), 1):
                if line.strip():
                    cells = line.split(',') if ',' in line else [line]
                    for j, cell in enumerate(cells, 1):
                        ws.cell(row=i, column=j, value=cell.strip())
            wb.save(str(file_path))
        else:
            with open(file_path, "w") as f:
                f.write(content)
        
        file_url = f"/api/static/files/{file_filename}"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "file", "prompt": request.prompt, "url": file_url,
            "file_type": ext, "content": content if not is_binary else f"[{ext.upper()} file]", "timestamp": timestamp
        })
        
        return {"id": gen_id, "file_url": file_url, "content": content if not is_binary else f"[{ext.upper()} file generated]", "file_type": ext, "timestamp": timestamp}
        
    except Exception as e:
        logger.error(f"File generation error: {e}")
        raise HTTPException(status_code=500, detail=f"File generation failed: {str(e)}")

# ============== DOCUMENT STUDIO ==============

@api_router.post("/document/generate")
async def generate_document(data: dict, user = Depends(get_current_user)):
    """GAAIUS AI Document Studio - Generate professional documents with AI Agents"""
    try:
        prompt = data.get("prompt", "")
        doc_type = data.get("document_type", "pdf")
        current_content = data.get("current_content", "")
        doc_name = data.get("document_name", "document")
        agent = data.get("agent", "general")
        
        # AI Agent personas with domain expertise
        agent_prompts = {
            "general": "You are GAAIUS AI, a professional document writer. Create well-structured, clear, and professional documents.",
            "lawyer": """You are GAAIUS AI Legal Counsel, an expert legal document writer. 
You specialize in contracts, agreements, NDAs, terms of service, privacy policies, and legal correspondence.
Always use precise legal language, include standard legal clauses (force majeure, indemnification, severability, governing law).
Add signature blocks and witness sections where appropriate. Include effective dates and term durations.
Note: This is for informational purposes. Recommend professional legal review for binding documents.""",
            "accountant": """You are GAAIUS AI Financial Expert, a certified accounting and finance professional.
You specialize in invoices, financial reports, budgets, tax documents, profit/loss statements, and balance sheets.
Always include proper calculations, tax breakdowns, payment terms, and currency formatting.
Use standard accounting formats (GAAP/IFRS style). Include totals, subtotals, and clear line items.""",
            "hr": """You are GAAIUS AI HR Manager, an expert in human resources documentation.
You specialize in job descriptions, employee handbooks, offer letters, performance reviews, termination letters, and company policies.
Use inclusive language, clear expectations, and standard HR formatting.
Include relevant sections for benefits, compensation, reporting structure, and compliance requirements.""",
            "marketing": """You are GAAIUS AI Marketing Director, an expert copywriter and strategist.
You specialize in business proposals, marketing plans, sales decks, press releases, case studies, and brand guidelines.
Write with persuasive, engaging language. Use data-driven arguments, clear value propositions, and compelling CTAs.
Include executive summaries, market analysis, and ROI projections where relevant.""",
            "academic": """You are GAAIUS AI Academic Writer, an expert in scholarly and research documentation.
You specialize in research papers, thesis outlines, literature reviews, grant proposals, and academic reports.
Use proper academic formatting (APA/MLA style), citations, abstract, methodology sections.
Include references section and maintain formal, objective academic tone throughout."""
        }
        
        agent_system = agent_prompts.get(agent, agent_prompts["general"])
        
        # Document type specific instructions
        doc_prompts = {
            "invoice": """You are a professional invoice generator. Create a detailed, professional invoice with:
- Company/Sender information (placeholder for user to fill)
- Client/Bill To information
- Invoice number and date
- Itemized list with descriptions, quantities, rates, amounts
- Subtotal, Tax (if applicable), Total
- Payment terms and bank details
- Professional formatting with clear sections""",
            
            "contract": """You are a legal document writer. Create a comprehensive contract/agreement with:
- Party information sections
- Detailed terms and conditions
- Scope of work/services
- Payment terms
- Duration and termination clauses
- Confidentiality clause
- Dispute resolution
- Signature blocks
Use professional legal language.""",
            
            "proposal": """You are a business proposal writer. Create a compelling business proposal with:
- Executive Summary
- Problem Statement
- Proposed Solution
- Methodology/Approach
- Timeline and Milestones
- Team/Qualifications
- Pricing/Investment
- Terms and Conditions
- Call to Action
Use persuasive, professional language.""",
            
            "resume": """You are a professional CV/resume writer. Create a modern, ATS-friendly resume with:
- Contact Information
- Professional Summary
- Skills section
- Work Experience (reverse chronological)
- Education
- Certifications/Awards
Use action verbs and quantifiable achievements.""",
            
            "report": """You are a professional report writer. Create a detailed report with:
- Executive Summary
- Introduction
- Methodology
- Findings/Results
- Analysis
- Conclusions
- Recommendations
- References
Use clear headings and professional formatting.""",
            
            "letter": """You are a professional letter writer. Create a well-formatted business letter with:
- Date
- Recipient information
- Subject line
- Salutation
- Body paragraphs
- Closing
- Signature block
Use appropriate formal tone.""",
            
            "xlsx": """You are a spreadsheet/data expert. Create structured data that works well in Excel:
- Use comma-separated values
- Include clear headers in first row
- Use proper data types (numbers, dates, text)
- Include calculations/formulas descriptions
- Organize data logically""",
            
            "default": """You are a professional document writer. Create well-structured content with:
- Clear headings using markdown (# ## ###)
- Organized sections
- Professional language
- Proper formatting"""
        }
        
        system_prompt = f"{agent_system}\n\n{doc_prompts.get(doc_type, doc_prompts['default'])}"
        
        # If editing existing content
        user_prompt = prompt
        if current_content:
            user_prompt = f"Current document content:\n{current_content[:2000]}\n\nUser request: {prompt}\n\nModify the document according to the request."
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.4,
            max_tokens=6000
        )
        
        content = completion.choices[0].message.content
        
        # Clean markdown code blocks
        if "```" in content:
            import re
            code_match = re.search(r'```[\w]*\n?([\s\S]*?)```', content)
            if code_match:
                content = code_match.group(1).strip()
        
        gen_id = str(uuid.uuid4())
        
        # Determine file extension
        ext_map = {
            "pdf": "pdf", "docx": "docx", "xlsx": "xlsx",
            "invoice": "pdf", "contract": "pdf", "proposal": "pdf",
            "resume": "pdf", "report": "pdf", "letter": "pdf",
            "presentation": "md"
        }
        ext = ext_map.get(doc_type, "md")
        
        file_filename = f"{gen_id}.{ext}"
        file_path = ROOT_DIR / "static" / "files" / file_filename
        (ROOT_DIR / "static" / "files").mkdir(parents=True, exist_ok=True)
        
        # Generate file based on type
        if ext == "pdf":
            try:
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                from reportlab.lib.units import inch
                
                doc = SimpleDocTemplate(str(file_path), pagesize=A4, 
                    leftMargin=0.75*inch, rightMargin=0.75*inch,
                    topMargin=0.75*inch, bottomMargin=0.75*inch)
                styles = getSampleStyleSheet()
                
                # Custom styles
                styles.add(ParagraphStyle(name='Title', parent=styles['Heading1'], fontSize=18, spaceAfter=20))
                styles.add(ParagraphStyle(name='Subtitle', parent=styles['Heading2'], fontSize=14, spaceAfter=12))
                
                story = []
                
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        story.append(Spacer(1, 6))
                    elif line.startswith('# '):
                        story.append(Paragraph(line[2:], styles['Title']))
                    elif line.startswith('## '):
                        story.append(Paragraph(line[3:], styles['Subtitle']))
                    elif line.startswith('### '):
                        story.append(Paragraph(line[4:], styles['Heading3']))
                    elif line.startswith('- ') or line.startswith('* '):
                        story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                    elif line.startswith(tuple('0123456789')):
                        story.append(Paragraph(line, styles['Normal']))
                    else:
                        story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 4))
                
                doc.build(story)
            except Exception as pdf_err:
                logger.error(f"PDF generation error: {pdf_err}")
                # Fallback to text file
                ext = "md"
                file_filename = f"{gen_id}.md"
                file_path = ROOT_DIR / "static" / "files" / file_filename
                with open(file_path, "w") as f:
                    f.write(content)
                    
        elif ext == "docx":
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                
                doc = Document()
                for line in content.split('\n'):
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.strip():
                        doc.add_paragraph(line)
                doc.save(str(file_path))
            except Exception as docx_err:
                logger.error(f"DOCX generation error: {docx_err}")
                ext = "md"
                file_filename = f"{gen_id}.md"
                file_path = ROOT_DIR / "static" / "files" / file_filename
                with open(file_path, "w") as f:
                    f.write(content)
                    
        elif ext == "xlsx":
            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment
                
                wb = Workbook()
                ws = wb.active
                ws.title = "Data"
                
                for i, line in enumerate(content.split('\n'), 1):
                    if line.strip():
                        cells = line.split(',') if ',' in line else line.split('\t') if '\t' in line else [line]
                        for j, cell in enumerate(cells, 1):
                            ws.cell(row=i, column=j, value=cell.strip())
                            if i == 1:  # Header row
                                ws.cell(row=i, column=j).font = Font(bold=True)
                wb.save(str(file_path))
            except Exception as xlsx_err:
                logger.error(f"XLSX generation error: {xlsx_err}")
                ext = "csv"
                file_filename = f"{gen_id}.csv"
                file_path = ROOT_DIR / "static" / "files" / file_filename
                with open(file_path, "w") as f:
                    f.write(content)
        else:
            with open(file_path, "w") as f:
                f.write(content)
        
        file_url = f"/api/static/files/{file_filename}"
        timestamp = datetime.now(timezone.utc).isoformat()
        
        await db.generations.insert_one({
            "id": gen_id, "type": "document", "prompt": prompt, "url": file_url,
            "document_type": doc_type, "content": content[:1000], "timestamp": timestamp
        })
        
        return {
            "id": gen_id, 
            "file_url": file_url, 
            "filename": f"{doc_name}.{ext}",
            "content": content,
            "document_type": doc_type,
            "message": f"Your {doc_type.upper()} document has been created! You can preview it and download.",
            "timestamp": timestamp
        }
        
    except Exception as e:
        logger.error(f"Document generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")

# ============== PROJECTS ==============

@api_router.post("/projects")
async def create_project(data: ProjectCreate, user = Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Login required")
    
    project = {
        "id": str(uuid.uuid4()),
        "user_id": user["id"],
        "name": data.name,
        "description": data.description,
        "type": data.type,
        "files": {},
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    await db.projects.insert_one(project)
    # Return without _id
    project.pop("_id", None)
    return project

@api_router.get("/projects")
async def get_projects(user = Depends(get_current_user)):
    if not user:
        return []
    projects = await db.projects.find({"user_id": user["id"]}, {"_id": 0}).sort("updated_at", -1).to_list(100)
    return projects

@api_router.get("/projects/{project_id}")
async def get_project(project_id: str, user = Depends(get_current_user)):
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

@api_router.put("/projects/{project_id}/files")
async def update_project_files(project_id: str, files: dict, user = Depends(get_current_user)):
    await db.projects.update_one(
        {"id": project_id},
        {"$set": {"files": files, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"status": "updated"}

# ============== BUILD/VIBE CODING ==============

@api_router.post("/build/generate")
async def build_generate(data: dict, user = Depends(get_current_user)):
    """Generate code for vibe coding builder"""
    try:
        prompt = data.get("prompt", "")
        current_code = data.get("current_code", "")
        
        system_prompt = """You are an expert web developer helping build React applications.
        Generate or modify code based on the user's request.
        Always output valid React/JSX code that can be rendered.
        Use Tailwind CSS for styling.
        Output ONLY the code, no explanations or markdown."""
        
        messages = [
            {"role": "system", "content": system_prompt},
        ]
        
        if current_code:
            messages.append({"role": "user", "content": f"Current code:\n```\n{current_code}\n```\n\nModify it to: {prompt}"})
        else:
            messages.append({"role": "user", "content": f"Create a React component: {prompt}"})
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.3,
            max_tokens=4096
        )
        
        code = completion.choices[0].message.content
        # Clean up code blocks if present
        if "```" in code:
            import re
            code_match = re.search(r'```(?:jsx?|tsx?|javascript|typescript)?\n?([\s\S]*?)```', code)
            if code_match:
                code = code_match.group(1)
        
        return {"code": code.strip(), "model_used": "Groq Llama 3.3"}
        
    except Exception as e:
        logger.error(f"Build generate error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/build/generate-full")
async def build_generate_full(data: dict, user = Depends(get_current_user)):
    """Generate a full web project with multiple files"""
    try:
        prompt = data.get("prompt", "")
        current_files = data.get("current_files", {})
        project_type = data.get("project_type", "web")
        
        system_prompt = """You are an expert full-stack web developer. You build REAL, functional websites and applications.

When the user asks you to build something, you must:
1. Create complete, working HTML files with embedded Tailwind CSS
2. Create proper JavaScript for interactivity
3. Create CSS for custom styling
4. Make it fully functional - not demos or mockups

Output format: Return a JSON object with:
- "files": an object where keys are filenames and values are the complete file contents
- "message": a brief description of what you built

Example response format:
{
  "files": {
    "index.html": "<!DOCTYPE html>...",
    "script.js": "// JavaScript code...",
    "style.css": "/* CSS styles */"
  },
  "message": "I built a responsive landing page with..."
}

IMPORTANT:
- Use Tailwind CSS via CDN in HTML
- Make the code production-ready
- Include proper meta tags and structure
- Add real functionality, not placeholder text
- Output ONLY valid JSON, no markdown or explanations"""
        
        # Build context from current files
        files_context = ""
        if current_files:
            files_context = "Current project files:\n"
            for filename, content in current_files.items():
                files_context += f"\n--- {filename} ---\n{content[:500]}...\n"
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"{files_context}\n\nUser request: {prompt}\n\nGenerate the updated/new files as JSON."}
        ]
        
        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.3,
            max_tokens=8000
        )
        
        response_text = completion.choices[0].message.content.strip()
        
        # Try to parse JSON response
        try:
            # Remove markdown code blocks if present
            if response_text.startswith("```"):
                response_text = response_text.split("```")[1]
                if response_text.startswith("json"):
                    response_text = response_text[4:]
            
            result = json.loads(response_text)
            return {
                "files": result.get("files", {}),
                "message": result.get("message", "Code updated!")
            }
        except json.JSONDecodeError:
            # If not valid JSON, treat as single HTML file update
            return {
                "files": {"index.html": response_text},
                "message": "I've updated your index.html"
            }
        
    except Exception as e:
        logger.error(f"Build generate-full error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ============== STATIC FILES ==============

@api_router.get("/static/{filename}")
async def serve_static(filename: str):
    file_path = ROOT_DIR / "static" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path)

@api_router.get("/static/videos/{filename}")
async def serve_video(filename: str):
    file_path = ROOT_DIR / "static" / "videos" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Video not found")
    return FileResponse(file_path, media_type="video/mp4")

@api_router.get("/static/audio/{filename}")
async def serve_audio(filename: str):
    file_path = ROOT_DIR / "static" / "audio" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Audio not found")
    return FileResponse(file_path, media_type="audio/wav")

@api_router.get("/static/files/{filename}")
async def serve_file(filename: str):
    file_path = ROOT_DIR / "static" / "files" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path)

# ============== GENERATIONS ==============

@api_router.get("/generations")
async def get_generations(gen_type: Optional[str] = None, limit: int = 20):
    query = {}
    if gen_type:
        query["type"] = gen_type
    generations = await db.generations.find(query, {"_id": 0}).sort("timestamp", -1).to_list(limit)
    return generations

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
